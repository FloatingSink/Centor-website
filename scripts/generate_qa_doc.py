from docx import Document
from docx.shared import Pt
import datetime

doc = Document()

# ── Helpers ──────────────────────────────────────────────────────────────────

def h(text, level):
    doc.add_heading(text, level=level)

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def numbered(text):
    doc.add_paragraph(text, style='List Number')

def verdict_para(label, text):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)

def issue_block(sev, en_src, zh_txt, issue_type, explanation):
    p = doc.add_paragraph()
    p.add_run("Severity: " + sev).bold = True
    doc.add_paragraph("  English source: " + en_src)
    doc.add_paragraph("  Chinese text: " + zh_txt)
    doc.add_paragraph("  Issue type: " + issue_type)
    doc.add_paragraph("  Explanation: " + explanation)
    doc.add_paragraph("")

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
h("CENTOR Global — Bilingual QA Review", 0)
para("English–Chinese Website Localisation Audit", bold=True)
today = datetime.date.today().strftime("%d %B %Y")
para("Review Date: " + today)
para("Website Reviewed: https://staging.centorglobal.com")
para("Scope: Engineering and technical terminology accuracy — English vs. Simplified Chinese")
doc.add_paragraph("")

# ══════════════════════════════════════════════════════════════════════════════
# 1. SITE-LEVEL SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h("1. Site-Level Summary", 1)

bullet("Total pages checked: 6 (index, about, products, services, references, contact)")
bullet("EN/ZH page pairs found: 6 confirmed pairs, mirrored under /zh/ subdirectory")
bullet("Pages with no clear language pair: None identified")
bullet(
    "Overall translation reliability: Generally functional. The bulk of the site translates "
    "accurately at a surface level. However, there are two high-severity factual errors in the "
    "Chinese homepage (wrong TBM type for a cited project; project names that do not match those "
    "on the references page), and one high-severity mission-statement mistranslation on the About "
    "page that changes the core technical claim."
)
bullet(
    "Engineering terminology reliability: Mixed. Standard tunneling terms (lip seal / 唇形密封, "
    "main bearing / 主轴承, tail seal / 盾尾, foam agent / 泡沫剂, penetration / 锥入度) are "
    "rendered correctly. Broader process and performance terms show weakness: 'uptime' becomes "
    "'efficiency,' 'soil conditioning' is rendered as 'muck improvement,' and 'Slurry TBM' is "
    "inconsistently labelled across pages. Product-tier terminology (Gold/Silver Line) is "
    "handled consistently."
)

h("Repeated Cross-Site Concerns", 2)
bullet(
    "'Soil conditioning' → '渣土改良' used site-wide. "
    "渣土 (excavated muck/spoil) is not the soil face being conditioned; "
    "correct term is 土体改良 or 地层改良."
)
bullet(
    "Slurry TBM rendered as '泥水加压 TBM' on one page and '泥水平衡 TBM' on another; "
    "neither matches the industry-standard '泥水盾构' or '泥水平衡盾构'."
)
bullet(
    "Project example names on the homepage (Chinese) do not match the names on the References page "
    "(Chinese or English), suggesting two separate sets of example projects were used without "
    "cross-checking."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. TERMINOLOGY GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
h("2. Terminology Glossary Check", 1)

glossary_rows = [
    ("Tunnel Boring Machine (TBM)", "隧道掘进机（TBM）", "Correct", "Standard Chinese equivalent"),
    ("EPB TBM", "EPB TBM", "Correct", "Acronym retained; acceptable in industry"),
    ("Slurry TBM", "泥水加压 TBM / 泥水平衡 TBM", "Inconsistent",
     "Two different renderings; standard is 泥水盾构 or 泥水平衡盾构"),
    ("Hard Rock TBM", "硬岩 TBM", "Correct", "Standard term"),
    ("Mixed-face TBM", "(omitted)", "Potentially incorrect",
     "复合地层 or 复合式 is the standard descriptor; omitted on references page"),
    ("Main bearing", "主轴承", "Correct", ""),
    ("Lip seal / lip-seal system", "唇形密封系统", "Correct", ""),
    ("Tail seal grease / tail sealant", "盾尾密封脂", "Correct", "Standard term"),
    ("Tail brush / tail sealing brush", "密封刷", "Correct", ""),
    ("Soil conditioning", "渣土改良", "Acceptable but imprecise",
     "渣土 = excavated spoil; 土体改良 or 地层改良 is more precise"),
    ("Foam agent", "泡沫剂", "Correct", "Standard term"),
    ("Clay-thickened", "黏土稠化", "Correct", ""),
    ("Penetration (NLGI grease)", "锥入度", "Correct", "Standard Chinese NLGI term"),
    ("Biodegradable", "可生物降解", "Correct", ""),
    ("Flame retardant", "阻燃", "Correct", ""),
    ("Self-extinguishing", "自熄性", "Correct", ""),
    ("Water seal (pressure rating)", "止水性能", "Correct", ""),
    ("Anti-washout", "抗冲刷", "Correct", ""),
    ("Expansion rate (foam)", "发泡倍率", "Correct", "Standard term"),
    ("Dispersive ground", "分散型地层", "Acceptable", "Context-dependent; acceptable for foam agent application"),
    ("Cohesive ground", "黏性地层", "Correct", ""),
    ("TBM uptime", "运转效率", "Potentially incorrect", "Uptime = availability / 稼动率, not efficiency / 效率"),
    ("Commissioning", "调试", "Correct", ""),
    ("Post-drive analysis", "掘进后分析", "Correct", ""),
    ("Advance rate", "掘进速度", "Correct", ""),
    ("Supply chain", "供应链", "Correct", ""),
    ("Gold Line / Silver Line", "金线 / 银线", "Correct", "Brand tiers translated by meaning consistently"),
    ("IBC (Intermediate Bulk Container)", "IBC", "Correct", "Acronym retained"),
    ("Technical data sheet (TDS)", "TDS", "Correct", "Acronym retained"),
    ("Safety data sheet (SDS)", "SDS", "Correct", "Acronym retained"),
]

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, txt in enumerate(["English Term", "Chinese Term Used", "Assessment", "Notes"]):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].bold = True

for row_data in glossary_rows:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_paragraph("")

# ══════════════════════════════════════════════════════════════════════════════
# 3. PAGE-BY-PAGE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
h("3. Page-by-Page Review", 1)

# ─── HOMEPAGE ────────────────────────────────────────────────────────────────
h("Page: Homepage", 2)
para("English URL: https://staging.centorglobal.com/index.html")
para("Chinese URL: https://staging.centorglobal.com/zh/index.html")
verdict_para("Verdict: ", "Potentially inaccurate technical translation")

h("Issues Found", 3)
issue_block(
    "High",
    "Project example — Dubai metro project",
    "迪拜地铁蓝线（地铁·泥水平衡 TBM）",
    "Mistranslation / wrong technical register",
    "The References page (both EN and ZH) identifies this project as the Dubai Metro Route 2020 "
    "Extension using EPB TBMs, not Slurry TBMs. The Chinese homepage names it the 'Blue Line' "
    "(蓝线) and assigns a Slurry TBM (泥水平衡). Two separate errors: wrong project name and wrong "
    "TBM type. Route 2020 Extension used EPB TBMs."
)
issue_block(
    "High",
    "Homepage project examples — Singapore, Hong Kong, Dubai",
    "ZH: 新加坡滨海汤东线 / 港铁屯马线延线 / 迪拜地铁蓝线",
    "Inconsistent term (cross-page)",
    "References page ZH names: 南北走廊隧道 / 屯门—赤鱲角连接路 / 迪拜地铁2020路线延伸. "
    "All three names differ from homepage Chinese. Technical readers familiar with these projects "
    "will notice immediately."
)
issue_block(
    "Medium",
    "\"Soil Conditioning — Foam Agents\"",
    "渣土改良 — 泡沫剂",
    "Overly generic / imprecise terminology",
    "渣土 refers to excavated spoil (muck) already removed. Soil conditioning is a process applied "
    "to the in-situ soil face and chamber. More precise term: 土体改良 or 地层改良. This issue "
    "recurs across all pages."
)
h("Consistency Notes", 3)
para(
    "'Slurry TBM' is rendered as 泥水平衡 TBM on the homepage but as 泥水加压 TBM on the "
    "References page. Neither matches the standard Chinese term 泥水盾构 / 泥水平衡盾构."
)
verdict_para("Needs human engineering review? ", "Yes")

# ─── ABOUT ───────────────────────────────────────────────────────────────────
h("Page: About Us", 2)
para("English URL: https://staging.centorglobal.com/about.html")
para("Chinese URL: https://staging.centorglobal.com/zh/about.html")
verdict_para("Verdict: ", "Generally accurate with minor terminology issues")

h("Issues Found", 3)
issue_block(
    "High",
    "Mission: \"Maximise TBM Uptime Through Superior Chemistry\"",
    "以卓越化学性能最大化TBM运转效率",
    "Mistranslation",
    "'Uptime' (正常运行时间 / 稼动率 / 运转率) is a measure of operational availability — "
    "the proportion of time a machine is running without downtime. '运转效率' means operational "
    "efficiency (how well it operates when running). These are distinct engineering metrics. For a "
    "company whose core value proposition is reducing TBM downtime, this distinction is "
    "commercially and technically significant."
)
issue_block(
    "Medium",
    "\"20+ years of combined underground project experience\"",
    "数十年实战经验",
    "Overly generic terminology",
    "'数十年' means 'several decades' and is ambiguous (could mean 20, 30, or 40+ years). "
    "The English specifies '20+' as a minimum benchmark. Loss of specificity weakens the claim."
)
issue_block(
    "Low",
    "\"Partnership\" (core value)",
    "深度合作",
    "Overly generic terminology",
    "Partnership in an engineering services context implies shared accountability and joint "
    "problem-solving. '深度合作' (deep cooperation) is weaker and more transactional. "
    "'伙伴关系' would be more precise."
)
verdict_para("Needs human engineering review? ", "Yes")

# ─── PRODUCTS ────────────────────────────────────────────────────────────────
h("Page: Products", 2)
para("English URL: https://staging.centorglobal.com/products.html")
para("Chinese URL: https://staging.centorglobal.com/zh/products.html")
verdict_para("Verdict: ", "Generally accurate with minor terminology issues")

h("Issues Found", 3)
issue_block(
    "Medium",
    "TM-T100: \"Economical driving-grade tail sealant for shielded TBMs\"",
    "防止水、土及环形注浆浆液进入 TBM 盾尾",
    "Omission / wrong technical register",
    "The Chinese substitutes the product classification descriptor (economical driving-grade) with "
    "a purely functional description (preventing ingress of water, soil, and annular grout). "
    "'经济型' and '掘进级' are product positioning terms that inform procurement decisions and "
    "are absent from the Chinese."
)
issue_block(
    "Medium",
    "PA — Tail Sealant Hand-coat: \"Water Seal: 35 bar\"",
    "止水性能：35 bar（松村标准）",
    "Addition not in English source",
    "The Chinese adds '（松村标准）' — a reference to the Matsumura test standard — which does "
    "not appear in the English specification. If this is a legitimate test protocol for the 35 bar "
    "rating, it should appear in both language versions. If not, it is an asymmetric technical "
    "claim that requires verification against the master product data sheet."
)
issue_block(
    "Low",
    "EP1/EP2: \"Eco-friendly sealant for the main bearing lip-seal system\"",
    "可生物降解、无毒——对地下水及土壤零污染",
    "Addition not in English source",
    "'Zero contamination to groundwater and soil' is a more specific environmental claim than "
    "'eco-friendly.' This elaboration is not reflected in the English source and should be "
    "verified before being treated as a product claim."
)
issue_block(
    "Low",
    "UK Gold Line project product codes: CTG-MB-GL01, CTS-ST-GL02 (on References page)",
    "Not mentioned in Chinese references",
    "Inconsistent term (cross-page)",
    "References page uses product codes CTG-MB-GL01 and CTS-ST-GL02, but the Products page uses "
    "PX2350, TM-T100, PA, EP1/EP2. Discrepancy is unexplained in either language."
)
verdict_para("Needs human engineering review? ", "Yes")

# ─── SERVICES ────────────────────────────────────────────────────────────────
h("Page: Services", 2)
para("English URL: https://staging.centorglobal.com/services.html")
para("Chinese URL: https://staging.centorglobal.com/zh/services.html")
verdict_para("Verdict: ", "Generally accurate with minor terminology issues")

h("Issues Found", 3)
issue_block(
    "Medium",
    "On-Site Technical Support: \"commissioning support, real-time monitoring, operator "
    "training, and troubleshooting\"",
    "TBM 调试、始发及地层过渡段支持",
    "Omission",
    "Chinese removes three service elements (real-time monitoring, operator training, "
    "troubleshooting) and adds two not in the English (TBM launch / 始发, geological transition "
    "zone support / 地层过渡段). This restructures the stated service scope and omits training "
    "and troubleshooting as offered deliverables."
)
issue_block(
    "Low",
    "Technical Documentation: \"compliance certificates\"",
    "竣工报告 (completion report)",
    "Mistranslation / wrong technical register",
    "A compliance certificate (合规证书) attests that a product or process meets a standard or "
    "regulation. A 竣工报告 is a project completion report. These are different document types "
    "serving different purposes."
)
issue_block(
    "Low",
    "Technical Documentation: \"safety documentation\"",
    "SDS only",
    "Overly specific / different scope",
    "'Safety documentation' is a broader category than SDS alone. Minor scope narrowing; "
    "likely acceptable in a product documentation context."
)
verdict_para("Needs human engineering review? ", "No")

# ─── REFERENCES ──────────────────────────────────────────────────────────────
h("Page: References", 2)
para("English URL: https://staging.centorglobal.com/references.html")
para("Chinese URL: https://staging.centorglobal.com/zh/references.html")
verdict_para("Verdict: ", "Generally accurate with minor terminology issues")

h("Issues Found", 3)
issue_block(
    "Medium",
    "Hong Kong: \"Slurry TBM Ø 14.0 m\"",
    "泥水加压 TBM Ø 14.0 m",
    "Overly generic terminology",
    "'泥水加压' (slurry pressurised) is a non-standard rendering. The established Chinese term is "
    "'泥水平衡盾构' or '泥水盾构'. On the homepage Chinese, the same machine type was called "
    "'泥水平衡 TBM' — inconsistency across two pages for the same TBM category."
)
issue_block(
    "Medium",
    "UK featured project: \"8 mixed-face EPB TBMs, Ø 7.1 m\"",
    "EPB TBM × 8 台，直径 Ø 7.1 m",
    "Omission",
    "'Mixed-face' (复合地层 / 复合地层适应型) indicates the TBM bores through heterogeneous "
    "ground with both soft and hard materials simultaneously — a higher-capability specification "
    "than a standard EPB TBM. Its omission understates the reference project's technical "
    "complexity for Chinese-speaking readers."
)
issue_block(
    "Low",
    "UK project: product codes CTG-MB-GL01, CTS-ST-GL02 listed in English",
    "Product codes omitted in Chinese version",
    "Omission",
    "Chinese version of the references page omits product codes cited in the UK Gold Line "
    "featured project, creating an asymmetry in technical detail available to Chinese-language "
    "readers."
)
h("Consistency Notes", 3)
para(
    "Project names on the References page (Chinese) are internally consistent. However, they do "
    "not match the project names shown on the homepage (Chinese). This is the most disruptive "
    "cross-page inconsistency on the site."
)
verdict_para("Needs human engineering review? ", "Yes")

# ─── CONTACT ─────────────────────────────────────────────────────────────────
h("Page: Contact", 2)
para("English URL: https://staging.centorglobal.com/contact.html")
para("Chinese URL: https://staging.centorglobal.com/zh/contact.html")
verdict_para("Verdict: ", "Generally accurate with minor terminology issues")

h("Issues Found", 3)
issue_block(
    "Medium",
    "Regional Offices: Dubai, London, Sydney, Taipei, Houston (5 named + Singapore HQ = 6 total locations)",
    "新加坡总部，全球六处区域办公室 (HQ + 6 regional = 7 total locations implied)",
    "Wrong unit or number",
    "English names 5 regional offices. Chinese states 6 regional offices. Either the English is "
    "missing one regional office name, or the Chinese count is incorrect. Factual discrepancy "
    "that affects credibility with Chinese-speaking clients."
)
verdict_para("Needs human engineering review? ", "No")

# ══════════════════════════════════════════════════════════════════════════════
# 4. HIGHEST-RISK ISSUES
# ══════════════════════════════════════════════════════════════════════════════
h("4. Highest-Risk Terminology Issues", 1)
para(
    "The following 10 issues represent the greatest risk of technical misrepresentation, "
    "client confusion, or credibility damage:"
)
doc.add_paragraph("")

top10 = [
    "[Homepage / HIGH] Dubai project Chinese: '泥水平衡 TBM' contradicts both the English source "
    "and the References page, which clearly identify this project as an EPB TBM drive. Wrong TBM "
    "type is a verifiable factual error visible to technical procurement readers.",

    "[About / HIGH] Mission statement: 'Maximise TBM Uptime' → '最大化TBM运转效率.' Uptime "
    "(availability) is not the same as efficiency. This is the company's primary value proposition "
    "and the mistranslation changes what the company claims to do for its customers.",

    "[Homepage → References / HIGH] Project names on the Chinese homepage (新加坡滨海汤东线, "
    "港铁屯马线延线, 迪拜地铁蓝线) do not match project names on the Chinese references page "
    "(南北走廊隧道, 屯门—赤鱲角连接路, 迪拜地铁2020路线延伸). Technical readers familiar with the "
    "industry will notice immediately.",

    "[Products / MEDIUM] PA product: Chinese adds '（松村标准）' as the test method behind the "
    "35 bar water seal rating. If this standard is not in the English-language product data sheet, "
    "the Chinese version makes an asymmetric technical claim. Requires verification against master "
    "product specs.",

    "[Site-wide / MEDIUM] 'Soil Conditioning' → '渣土改良' used across the entire site. 渣土 "
    "(excavated muck) is the wrong reference material. The conditioning process acts on the in-situ "
    "soil/formation. Correct term: 土体改良 or 地层改良.",

    "[References / MEDIUM] 'Slurry TBM' → '泥水加压 TBM' (References page) and '泥水平衡 TBM' "
    "(Homepage Chinese). Two different Chinese terms for the same machine type across two pages. "
    "Standard industry term '泥水盾构' / '泥水平衡盾构' is used by neither.",

    "[Contact / MEDIUM] Regional office count: English lists 5, Chinese states 6. A Chinese-"
    "speaking client calling to locate a regional office may receive wrong expectations about "
    "CENTOR's geographic footprint.",

    "[References / MEDIUM] 'Mixed-face EPB TBM' for the UK featured project omitted in Chinese. "
    "Mixed-face capability is a key technical specification that differentiates this machine and "
    "project from standard EPB applications. Omission understates the reference project's "
    "complexity for Chinese-speaking readers.",

    "[Products / MEDIUM] TM-T100: product classification as 'economical driving-grade' omitted in "
    "Chinese, replaced with a functional description. Product-grade descriptors matter in "
    "procurement — '掘进级' indicates a performance tier, not merely a function.",

    "[Services / MEDIUM] 'On-Site Technical Support' scope restructured in Chinese: training and "
    "troubleshooting are removed and replaced with TBM launch and geological transition zone "
    "support. This changes the stated service scope and may affect what Chinese-speaking clients "
    "expect CENTOR to provide on site.",
]

for item in top10:
    numbered(item)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = (
    r"C:\Users\Jia Long Yu\OneDrive\Documents\GitHub\Centor-website"
    r"\CENTOR_Bilingual_QA_Review.docx"
)
doc.save(output_path)
print("Saved: " + output_path)
